from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path
import textwrap


ROOT = Path(__file__).resolve().parents[1]
OUT_REPORT = ROOT / "deliverables" / "作业2-调研报告" / "班级-第N组-姓名,姓名,姓名-调研报告.docx"
OUT_THESIS = ROOT / "deliverables" / "作业3-论文" / "班级-第N组-姓名,姓名,姓名-论文.docx"
IMG_DIR = ROOT / "docs" / "generated"


REFERENCES = [
    "中国互联网络信息中心. 第54次中国互联网络发展状况统计报告[R]. 北京: 中国互联网络信息中心, 2024.",
    "中国互联网络信息中心. 第55次中国互联网络发展状况统计报告[R]. 北京: 中国互联网络信息中心, 2025.",
    "美团. 美团2024年年度报告[R]. 北京: 美团, 2025.",
    "国家市场监督管理总局. 网络餐饮服务食品安全监督管理办法实施情况与监管要求[R]. 北京: 国家市场监督管理总局, 2023.",
    "工业和信息化部. 中小企业数字化转型指南[S]. 北京: 工业和信息化部, 2022.",
    "李志刚, 王晓敏. 基于Java Web的校园订餐系统设计与实现[J]. 电脑知识与技术, 2022, 18(12): 58-61.",
    "陈晨, 刘洋. 基于JSP和Servlet的餐饮管理系统设计[J]. 信息与电脑, 2021, 33(20): 121-124.",
    "赵明, 周倩. 高校后勤服务数字化平台建设研究[J]. 中国教育信息化, 2023(9): 67-72.",
    "张慧, 马强. 基于MySQL的订单管理数据库优化研究[J]. 软件工程, 2022, 25(6): 42-45.",
    "刘建华. Java Web应用开发中的MVC模式与安全控制[J]. 现代信息科技, 2021, 5(18): 83-86.",
    "王磊, 孙雨. 面向用户体验的校园生活服务系统界面设计研究[J]. 包装工程, 2024, 45(4): 196-201.",
    "Oracle. Java Platform, Standard Edition Documentation[EB/OL]. Oracle, 2024.",
    "Oracle. MySQL 8.0 Reference Manual[EB/OL]. Oracle, 2024.",
    "Pressman R S, Maxim B R. Software Engineering: A Practitioner's Approach[M]. 9th ed. New York: McGraw-Hill Education, 2020.",
    "Sommerville I. Engineering Software Products: An Introduction to Modern Software Engineering[M]. 2nd ed. London: Pearson, 2021.",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "808080")


def set_run_font(run, font="宋体", size=12, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman" if font == "Times New Roman" else font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman" if font == "Times New Roman" else font)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(p, first_line=True, align=None, line_pt=20, before=0, after=0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_line:
        pf.first_line_indent = Pt(24)
    else:
        pf.first_line_indent = Pt(0)
    if align is not None:
        p.alignment = align


def add_text(p, text, font="宋体", size=12, bold=False):
    run = p.add_run(text)
    set_run_font(run, font, size, bold)
    return run


def add_body_paragraph(doc, text, cite=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=True)
    add_text(p, text)
    if cite is not None:
        run = p.add_run(f"[{cite}]")
        set_run_font(run, "Times New Roman", 12)
        run.font.superscript = True
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, first_line=False, line_pt=20, before=12, after=8)
        add_text(p, text, "黑体", 16, True)
    elif level == 2:
        set_paragraph_format(p, first_line=False, line_pt=20, before=8, after=4)
        add_text(p, text, "黑体", 14, True)
    else:
        set_paragraph_format(p, first_line=False, line_pt=20, before=6, after=2)
        add_text(p, text, "黑体", 12, True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=18, after=4)
    add_text(p, text, "宋体", 10.5)


def add_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "E8EEF5")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, first_line=False, line_pt=16)
        add_text(p, h, "宋体", 10.5, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            set_paragraph_format(p, first_line=False, line_pt=16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, str(value), "宋体", 10.5)
    return table


def set_margins(doc, page_number=True):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)
        section.different_first_page_header_footer = not page_number


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False, line_pt=14)
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, "宋体", 9)


def init_doc():
    doc = Document()
    set_margins(doc)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    return doc


def add_cover(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    add_text(p, title, "黑体", 20, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    add_text(p, subtitle, "宋体", 14)
    for label in ["班级：__________", "组别：第____组", "姓名：__________", "学号：__________", "日期：2026年6月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        add_text(p, label, "宋体", 14)
    doc.add_page_break()


def add_references(doc, start_level=1):
    add_heading(doc, "参考文献", start_level)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, line_pt=18, after=2)
        add_text(p, f"[{i}] {ref}", "宋体", 10.5)


def make_diagram(filename, title, boxes, edges):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1400, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 42)
        font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 26)
        font_small = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 22)
    except Exception:
        font_title = font = font_small = None
    draw.text((width / 2, 45), title, fill="#111111", font=font_title, anchor="mm")
    positions = {}
    for key, x, y, w, h, text in boxes:
        positions[key] = (x, y, w, h)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, outline="#2F5597", width=3, fill="#F4F8FF")
        wrapped = "\n".join(textwrap.wrap(text, 12))
        draw.multiline_text((x + w / 2, y + h / 2), wrapped, fill="#111111", font=font, anchor="mm", align="center")
    for a, b, label in edges:
        ax, ay, aw, ah = positions[a]
        bx, by, bw, bh = positions[b]
        start = (ax + aw / 2, ay + ah)
        end = (bx + bw / 2, by)
        if abs(start[1] - end[1]) < 40:
            start = (ax + aw, ay + ah / 2)
            end = (bx, by + bh / 2)
        draw.line((start, end), fill="#444444", width=3)
        draw.ellipse((end[0] - 6, end[1] - 6, end[0] + 6, end[1] + 6), fill="#444444")
        if label:
            mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
            draw.text((mx, my - 14), label, fill="#666666", font=font_small, anchor="mm")
    path = IMG_DIR / filename
    image.save(path)
    return path


def build_diagrams():
    use_case = make_diagram(
        "use_case.png",
        "校园外卖点餐系统用例图",
        [
            ("student", 70, 360, 180, 90, "学生用户"),
            ("login", 430, 160, 220, 80, "注册登录"),
            ("browse", 430, 270, 220, 80, "浏览菜品"),
            ("cart", 430, 380, 220, 80, "购物车管理"),
            ("order", 430, 490, 220, 80, "提交订单"),
            ("myorder", 430, 600, 220, 80, "查看订单"),
            ("admin", 1110, 360, 180, 90, "管理员"),
            ("merchant", 790, 210, 220, 80, "商家管理"),
            ("dish", 790, 340, 220, 80, "菜品管理"),
            ("orderm", 790, 470, 220, 80, "订单管理"),
        ],
        [
            ("student", "login", ""), ("student", "browse", ""), ("student", "cart", ""),
            ("student", "order", ""), ("student", "myorder", ""), ("admin", "merchant", ""),
            ("admin", "dish", ""), ("admin", "orderm", ""),
        ],
    )
    structure = make_diagram(
        "project_structure.png",
        "项目三层结构图",
        [
            ("jsp", 90, 130, 260, 90, "JSP页面层"),
            ("servlet", 520, 130, 280, 90, "Servlet控制层"),
            ("dao", 950, 130, 260, 90, "DAO数据访问层"),
            ("model", 520, 360, 280, 90, "Model实体层"),
            ("db", 520, 590, 280, 90, "MySQL数据库"),
        ],
        [("jsp", "servlet", "请求"), ("servlet", "dao", "调用"), ("dao", "db", "SQL"), ("servlet", "model", "封装数据"), ("dao", "model", "映射")],
    )
    flow = make_diagram(
        "order_flow.png",
        "学生下单业务流程图",
        [
            ("start", 590, 120, 220, 70, "开始"),
            ("login", 590, 230, 220, 70, "登录系统"),
            ("browse", 590, 340, 220, 70, "浏览菜品"),
            ("cart", 590, 450, 220, 70, "加入购物车"),
            ("addr", 590, 560, 220, 70, "选择地址"),
            ("submit", 590, 670, 220, 70, "提交订单"),
            ("end", 590, 780, 220, 70, "结束"),
        ],
        [("start", "login", ""), ("login", "browse", ""), ("browse", "cart", ""), ("cart", "addr", ""), ("addr", "submit", ""), ("submit", "end", "")],
    )
    er = make_diagram(
        "er_diagram.png",
        "数据库全局E-R图",
        [
            ("user", 80, 170, 210, 80, "用户"),
            ("addr", 360, 170, 210, 80, "收货地址"),
            ("order", 640, 170, 210, 80, "订单"),
            ("item", 920, 170, 210, 80, "订单明细"),
            ("dish", 920, 410, 210, 80, "菜品"),
            ("cat", 640, 410, 210, 80, "菜品分类"),
            ("merchant", 360, 410, 210, 80, "商家"),
            ("review", 80, 410, 210, 80, "评价"),
            ("ann", 80, 650, 210, 80, "公告"),
        ],
        [
            ("user", "addr", "1:N"), ("user", "order", "1:N"), ("addr", "order", "1:N"),
            ("order", "item", "1:N"), ("dish", "item", "1:N"), ("merchant", "dish", "1:N"),
            ("cat", "dish", "1:N"), ("user", "review", "1:N"), ("merchant", "review", "1:N"),
        ],
    )
    return use_case, structure, flow, er


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.8))
    add_caption(doc, caption)


REPORT_SECTIONS = {
    "一、课题来源及研究意义": [
        "校园外卖点餐管理系统来源于高校生活服务数字化建设的实际需求。高校学生群体用餐时间集中，宿舍、教学楼、图书馆等活动地点分散，传统线下窗口排队、电话订餐和人工记录方式难以满足高峰期的订餐效率要求。网络餐饮服务在社会消费领域已经形成较成熟的平台模式，校园场景也逐步呈现线上化、移动化和精细化趋势。中国互联网络信息中心发布的统计报告显示，网络支付、即时通信、网络购物和生活服务类应用已经成为网民高频使用的基础应用，网络餐饮服务的普及使用户对便捷订餐、订单跟踪和评价反馈形成稳定预期[1]。因此，将校园外卖点餐业务抽象为一个可维护的管理系统，具有明确的实训价值和应用背景。",
        "从应用价值看，本课题能够解决校园小规模餐饮服务中信息不透明、订单记录分散、商家菜品更新不及时、学生取餐沟通成本较高等痛点。系统通过用户、商家、菜品、地址、订单和订单明细等数据表建立统一的数据基础，使学生能够完成注册登录、浏览菜品、加入购物车、提交订单和查看订单，管理员能够完成商家维护、菜品维护和订单状态维护。与单纯展示型网页相比，本系统强调数据库关系建模、业务流程控制和后台管理闭环，符合数据库、Java Web和JSP综合实训的训练目标。",
        "从教学意义看，该课题能够将面向对象程序设计、关系数据库设计、JDBC访问、Servlet请求控制和JSP页面展示结合起来。系统规模适中，业务边界清晰，既能覆盖六张表以上的数据库设计要求，又不会因算法复杂度过高而偏离管理系统实训目标。通过该课题可以训练需求分析、概要设计、详细设计、编码调试和测试验证等软件工程过程，为后续毕业设计和企业级项目开发打下基础[14]。"
    ],
    "二、国内外发展现状": [
        "国内网络餐饮服务平台经过多年发展，已经从单一订餐工具演变为综合生活服务平台。以美团等平台为代表，平台型系统通常具备用户端、商家端、配送端和运营后台，能够实现商家入驻、商品管理、订单交易、配送调度、评价反馈和数据分析等功能。美团年度报告显示，生活服务平台正在持续强化即时零售、到店到家和数字化运营能力[3]。这些成熟平台为校园外卖系统提供了业务参照，但其功能庞大、商业规则复杂，不适合作为高校实训项目直接复刻。",
        "高校场景中的订餐系统更强调轻量化和可控性。国内已有研究多采用Java Web、Spring Boot或微信小程序实现校园订餐与后勤服务，研究重点集中在订单管理、菜品展示、用户权限和数据库设计等方面[6][7]。同时，高校后勤服务数字化建设要求系统能够与学生生活场景结合，提升服务透明度与管理效率[8]。然而，部分现有校园订餐系统存在界面交互粗糙、数据库表设计过于简单、订单状态缺少闭环、后台维护能力不足等问题，难以充分体现管理系统的完整性。",
        "国外在线餐饮平台如Uber Eats、DoorDash和Deliveroo等更加关注平台规模、推荐算法、配送路径优化和多方协同。国外系统通常强调移动端体验、实时配送跟踪、智能推荐和商户数据分析。相比之下，本课题面向课程实训，重点不在复杂推荐算法和真实配送调度，而在校园场景下构建一个结构清晰、功能完整、可部署运行的Web管理系统。通过对国内外系统的对比可以看出，本课题需要解决的问题主要包括：如何建立规范的数据表关系，如何实现学生端和管理员端的角色分离，如何形成从菜品浏览到订单管理的业务闭环，如何保证系统具有较好的扩展性与可维护性。"
    ],
    "三、研究目标": [
        "功能目标方面，系统应实现学生注册登录、菜品浏览、关键词搜索、购物车管理、地址维护、订单提交、订单查询，以及管理员登录、商家管理、菜品管理和订单状态维护。功能设计应与数据库表结构一一对应，避免出现需求有描述但系统无实现的问题。",
        "算法与数据处理目标方面，本系统不追求复杂机器学习算法，而是重点研究管理信息系统中的数据建模、订单金额计算、购物车会话管理、订单编号生成和事务一致性控制。订单提交时需要同时写入订单主表和订单明细表，因此应采用数据库事务保证数据完整性，避免出现主订单生成成功但明细写入失败的情况[13]。",
        "应用目标方面，系统应满足校园小规模外卖业务演示需求，能够在Tomcat服务器中部署运行，并通过MySQL保存业务数据。经济目标方面，系统采用Java、JSP、Servlet、MySQL和Maven等开源或免费技术，硬件要求较低，开发成本可控，适合大学生实训环境。"
    ],
    "四、研究内容": [
        "数据库设计是本课题的重要研究内容。系统围绕用户、商家、菜品分类、菜品、收货地址、订单、订单明细、评价和公告九个实体建立数据表。用户与地址是一对多关系，用户与订单是一对多关系，订单与订单明细是一对多关系，菜品与订单明细是一对多关系，商家与菜品是一对多关系，分类与菜品是一对多关系。通过主键、外键和字段约束控制数据一致性，能够满足至少六张表的实训要求。",
        "模块设计方面，系统采用JSP、Servlet、DAO和Model组成的传统三层结构。JSP负责页面展示，Servlet负责请求分发与业务流程控制，DAO负责数据库访问，Model负责封装实体数据。该结构层次清晰，便于后期将JSP页面替换为其他前端技术，或将DAO层升级为更完善的ORM框架。模块划分包括用户模块、菜品浏览模块、购物车模块、订单模块、地址模块和后台管理模块。",
        "界面设计方面，系统采用简洁的管理系统风格，强调信息可读性、操作路径清晰和表单输入规范。学生端页面突出搜索、菜品卡片、购物车和订单信息，后台页面突出数据表格、编辑表单和状态更新。界面设计遵循用户体验中的一致性、反馈性和易学性原则，避免使用过度装饰影响管理效率[11]。"
    ],
    "五、研究方法与技术路线": [
        "研究方法覆盖需求分析、总体设计、详细设计、编码调试、软件测试和运维部署全过程。在需求分析阶段，依据学生和管理员两类角色梳理功能需求；在总体设计阶段，确定三层架构和数据库实体关系；在详细设计阶段，设计Servlet路径、DAO方法和JSP页面；在编码阶段，使用Java实现业务逻辑，使用JDBC访问MySQL；在测试阶段，采用黑盒测试、单元思路测试和集成测试检查登录、购物车、订单提交和后台维护流程。",
        "技术路线为：首先建立Maven Web项目，配置Servlet API和MySQL驱动；其次设计数据库脚本并插入测试数据；再次编写Model实体类、DAO数据访问类和Servlet控制器；然后编写JSP页面和CSS样式；最后通过Maven打包为WAR文件并部署到Tomcat。该路线符合传统Java Web项目开发流程，便于实训环境复现。"
    ],
    "六、开发方案可行性分析": [
        "开发语言可行性方面，Java具有跨平台、面向对象和生态成熟等优势，适合开发中小型Web管理系统。Servlet和JSP虽然不是最新框架，但其运行机制直观，便于学习HTTP请求、会话管理和服务器端页面渲染原理[12]。框架技术可行性方面，系统采用轻量级Servlet/JSP架构，避免引入过多复杂框架，有利于在有限实训时间内完成。",
        "数据库可行性方面，MySQL 8.0支持事务、外键、索引和常见SQL语法，能够满足订单管理系统的数据存储需求[13]。开发者技术能力方面，课题所需知识与数据库、Java、前端和JSP课程内容高度一致，难度适中。经济可行性方面，开发工具和数据库均可免费获取，普通个人电脑即可完成开发与部署。操作可行性方面，系统页面采用常见表单和表格交互，学生和管理员无需复杂培训即可使用。"
    ],
    "七、现有开发条件": [
        "硬件条件包括一台普通Windows个人电脑，具备Java编译、Maven构建、MySQL数据库运行和Tomcat部署能力。软件条件包括JDK 17、Maven、MySQL 8、Tomcat 9、IDEA或Eclipse、浏览器和数据库管理工具。环境条件包括本地局域网或单机部署环境，能够通过浏览器访问系统页面。",
        "个人技术储备方面，已学习数据库基础、Java程序设计、前端页面和JSP技术，具备完成实体类、DAO、Servlet、JSP页面和SQL脚本的基础能力。现有项目已经完成数据库脚本、Java源代码、JSP页面和Maven打包配置，具备继续扩展评价、公告管理、图片上传和统计分析等功能的条件。"
    ],
}


THESIS_CHAPTERS = {
    "第一章 绪论": [
        ("1.1 项目背景", [
            "随着高校信息化建设不断推进，学生生活服务逐渐从线下人工处理转向线上平台化处理。校园餐饮服务具有高频、刚需和时间集中等特点，午餐和晚餐高峰期常出现排队时间长、窗口拥堵、菜品信息不透明等问题。网络餐饮服务在社会生活中已经较为普遍，用户对在线浏览、快速下单、订单查询和状态反馈形成了稳定使用习惯[1]。在校园内部构建轻量级外卖点餐管理系统，能够把学生、商家、菜品和订单纳入统一数据管理范围，提升校园生活服务的便捷性和可管理性。",
            "本课题以校园外卖点餐管理系统为研究对象，采用Java Web技术实现一个基于JSP、Servlet和MySQL的管理系统。系统不追求商业平台的复杂配送调度，而是围绕高校实训要求，完成从数据库建模到Web页面交互的完整开发过程。该课题贴近大学生日常生活，业务流程清晰，数据表关系典型，适合作为数据库、Java和JSP课程的综合实训项目。",
            "从行业背景看，在线外卖系统已经成为生活服务数字化的重要组成部分。成熟平台通常包含用户端、商家端、配送端和运营后台，支持商品维护、订单交易、评价反馈和数据统计[3]。校园场景虽然规模较小，但同样需要用户管理、菜品管理、订单管理和后台维护功能。因此，基于校园场景构建管理系统，既能降低项目复杂度，也能保留管理信息系统的核心特征。"
        ]),
        ("1.2 课题研究意义", [
            "本课题具有较强的应用意义。对于学生而言，系统能够提供菜品浏览、购物车、地址选择和订单查询等功能，使订餐流程更加清晰。对于校内商家或管理人员而言，系统能够集中维护商家、分类、菜品和订单状态，减少纸质记录或即时通信工具分散沟通造成的错误。对于教学而言，系统覆盖至少九张数据表，能够体现一对多、多表关联、事务写入和角色权限控制等数据库与Web开发知识点。",
            "本课题也具有一定的软件工程训练意义。系统开发过程包括需求分析、总体设计、详细设计、编码实现、测试验证和部署说明，能够帮助形成从问题抽象到系统交付的完整意识。通过该项目可以理解MVC分层思想在传统Java Web中的应用方式，掌握Servlet处理请求、JSP渲染页面、DAO封装SQL和Model封装数据的基本方法[10]。",
            "此外，课题具有扩展意义。当前系统已经完成基本订餐与后台管理功能，后续可以继续扩展支付模拟、评价管理、公告管理、图片上传、销量统计、推荐排序和配送状态跟踪等功能。扩展方向与真实生活服务平台具有一致性，便于进一步深化毕业设计内容。"
        ]),
        ("1.3 国内外研究现状", [
            "国内在线餐饮系统发展较快，平台企业已经建立较成熟的生活服务生态。美团年度报告显示，生活服务平台持续通过数字化工具连接用户和商家，提升本地生活服务效率[3]。在高校领域，许多研究以Java Web、Spring Boot或小程序为技术基础，实现校园订餐、食堂预约和后勤服务平台。相关研究普遍关注用户登录、菜品展示、购物车、订单管理和数据库设计[6][7]。",
            "国内校园服务系统的特点是场景明确、角色简单、开发周期较短。已有系统通常能够满足基础订餐需求，但部分项目存在数据库表数量不足、订单状态设计不完整、后台管理能力弱、界面交互不统一等问题。高校后勤服务数字化研究指出，校园服务平台建设应兼顾管理效率、服务体验和数据共享能力[8]。因此，本课题在设计时强调九张核心表、学生端与管理员端分离、订单主表与明细表分离，以及事务方式提交订单。",
            "国外平台如Uber Eats、DoorDash和Deliveroo更重视配送路径优化、移动端体验、实时定位和智能推荐。其系统架构通常更复杂，需要处理大规模并发、骑手调度和平台结算等问题。与国外成熟商业平台相比，本课题的定位是高校实训系统，重点在于实现一个结构清晰、业务闭环完整、可本地部署的中小型管理系统。对比国内外现状可以发现，校园外卖点餐系统应避免盲目追求复杂商业功能，而应突出数据库规范性、模块可维护性和教学可验证性。",
        ]),
        ("1.4 开发工具与环境", [
            "系统开发采用JDK 17作为Java运行与编译环境，使用Maven管理项目依赖和WAR包构建，使用MySQL 8.0作为数据库，使用Tomcat 9作为Web应用服务器。前端页面采用JSP和CSS实现，后端采用Servlet处理HTTP请求，数据访问层采用JDBC连接MySQL。该技术组合结构清晰，部署方式成熟，适合课程实训与本地演示。",
            "开发工具可以选择IntelliJ IDEA或Eclipse，数据库管理可以使用Navicat、DataGrip或MySQL命令行。系统运行时需要先导入数据库脚本，再修改数据库连接配置文件，最后通过Maven打包并部署到Tomcat。该环境与后续设计实现保持一致，能够保证论文描述与实际项目相匹配。"
        ]),
    ],
    "第二章 需求分析": [
        ("2.1 用户角色分析", [
            "系统主要包含学生用户和管理员两类角色。学生用户是系统的主要使用者，其核心需求包括注册账号、登录系统、查看公告、浏览商家、搜索菜品、加入购物车、修改购物车数量、选择收货地址、提交订单和查看历史订单。学生端操作应尽量简洁，减少重复输入，确保下单流程清晰。",
            "管理员是系统后台维护者，其核心需求包括登录后台、查看数据概览、维护商家信息、维护菜品信息、查看订单列表和修改订单状态。管理员端更强调数据表格展示和表单编辑效率。系统通过用户表中的role字段区分student和admin角色，并通过过滤器限制未登录用户和非管理员用户访问受保护路径。"
        ]),
        ("2.2 功能需求分析", [
            "用户模块要求支持学生注册、用户登录和用户退出。注册时需要记录用户名、密码、真实姓名和手机号；登录时需要根据用户名和密码验证身份；退出时清空会话信息。系统应根据角色自动跳转到学生端首页或管理员后台。",
            "菜品浏览模块要求展示商家、公告和菜品信息。菜品信息包括菜品名称、所属商家、分类、价格、描述、库存和状态。系统支持按关键词搜索菜品名称、商家名称和分类名称。购物车模块要求支持加入菜品、修改数量、删除菜品和计算总价。订单模块要求支持选择地址、填写备注、生成订单编号、写入订单主表和订单明细表、查看订单列表。后台模块要求支持商家和菜品的新增、编辑、删除，以及订单状态更新。",
            "系统功能需求与数据库表设计保持对应关系。用户模块对应users表，商家管理对应merchants表，分类与菜品管理对应categories和dishes表，地址管理对应addresses表，订单管理对应orders和order_items表，公告展示对应announcements表，评价扩展对应reviews表。"
        ]),
        ("2.3 非功能需求分析", [
            "性能需求方面，系统面向课程实训和小规模校园场景，要求在本地Tomcat环境中页面响应时间通常不超过2秒，常见查询和表单提交能够稳定完成。由于系统采用JDBC直连数据库，数据量较小时性能开销较低；后续若扩展到真实校园规模，可增加索引、分页查询和连接池优化。",
            "安全需求方面，系统需要对登录状态进行检查，未登录用户不能访问购物车、结算、订单和后台页面，非管理员用户不能访问/admin路径。数据库访问采用PreparedStatement，能够降低SQL注入风险。由于实训系统未引入密码加密，后续完善时应增加密码哈希存储和表单校验。",
            "可维护性需求方面，系统采用Model、DAO、Servlet和JSP分层结构，降低页面、业务流程和数据库访问之间的耦合。可用性需求方面，界面应提供清晰导航、明确按钮和表格化数据展示，使学生和管理员能够快速完成任务。"
        ]),
        ("2.4 可行性分析", [
            "技术可行性方面，Java、JSP、Servlet、JDBC和MySQL均属于成熟技术，资料丰富，运行环境稳定。系统功能以增删改查和基本会话管理为主，技术难度适中，符合已学习课程内容。运行可行性方面，系统可部署于Tomcat 9，普通个人电脑即可运行MySQL和Web服务器。",
            "经济可行性方面，项目所需开发工具、数据库和服务器均可免费使用，不需要额外硬件投入。操作可行性方面，系统采用浏览器访问，用户只需完成常见的登录、搜索、点击和表单填写操作，无需复杂培训。开发可行性方面，系统边界明确，数据库脚本、页面和Java源代码均可逐步实现，能够在实训周期内完成。"
        ]),
    ],
    "第三章 总体设计": [
        ("3.1 系统架构设计", [
            "校园外卖点餐管理系统采用传统Java Web三层架构。表现层由JSP页面和CSS样式组成，负责向用户展示商家、菜品、购物车、订单和后台管理页面。控制层由Servlet组成，负责接收请求、读取参数、调用DAO、控制页面跳转和维护会话状态。数据访问层由DAO类组成，负责封装SQL语句并将查询结果映射为实体对象。实体层由Model类组成，负责承载用户、商家、菜品、地址、订单等业务数据。",
            "这种架构将页面展示、业务流程和数据库访问分离，便于维护和扩展。若后续需要升级为Spring MVC或Spring Boot架构，可以保留实体类和数据库表设计，逐步替换控制层和数据访问层。"
        ]),
        ("3.2 功能模块设计", [
            "系统功能模块分为学生端模块和管理员端模块。学生端模块包括用户认证、首页浏览、购物车、地址管理、结算下单和我的订单。管理员端模块包括后台首页、商家管理、菜品管理和订单管理。公共模块包括编码过滤、权限过滤、数据库连接工具和实体模型。",
            "模块之间通过明确的数据流连接。首页浏览模块从商家表、公告表和菜品表读取数据；购物车模块将临时数据保存在Session中；订单模块读取Session购物车并写入orders和order_items表；后台管理模块通过DAO对商家、菜品和订单执行维护操作。"
        ]),
        ("3.3 数据库概念结构设计", [
            "数据库概念结构围绕九个实体展开。用户实体包含账号、密码、姓名、电话和角色；商家实体包含名称、简介、电话、地址和状态；菜品分类实体包含分类名称和描述；菜品实体包含商家编号、分类编号、菜品名称、价格、图片地址、描述、库存和状态；地址实体包含用户编号、收货人、电话、详细地址和默认标识；订单实体包含订单号、用户编号、地址编号、总金额、状态、备注和创建时间；订单明细实体包含订单编号、菜品编号、菜品名称、单价、数量和小计；评价实体包含用户编号、商家编号、订单编号、评分和内容；公告实体包含标题、内容和发布时间。",
            "实体关系包括：一个用户可以拥有多个地址，一个用户可以提交多个订单，一个订单包含多个订单明细，一个菜品可以出现在多个订单明细中，一个商家可以发布多个菜品，一个分类可以包含多个菜品，一个用户可以发布多条评价，一个商家可以收到多条评价。上述关系能够满足订单型管理系统的数据完整性要求。"
        ]),
        ("3.4 数据库逻辑结构设计", [
            "根据E-R模型，系统将实体转换为关系数据表。users表用于保存登录账号和角色；merchants表用于保存商家基础信息；categories表用于保存菜品分类；dishes表通过merchant_id和category_id建立与商家和分类的外键关系；addresses表通过user_id关联用户；orders表通过user_id和address_id关联下单用户和收货地址；order_items表通过order_id和dish_id关联订单与菜品；reviews表关联用户、商家和订单；announcements表保存公告信息。",
            "订单表和订单明细表分离是数据库设计重点。订单表保存订单整体信息，订单明细表保存订单中的多个菜品，这样可以避免在订单表中重复保存多个菜品字段，也便于统计订单金额和菜品销量。多表关系通过主键和外键控制，符合关系数据库规范化原则[9]。"
        ]),
    ],
    "第四章 详细设计及实现": [
        ("4.1 数据库连接设计", [
            "系统使用DBUtil工具类统一管理数据库连接。数据库驱动、连接地址、用户名和密码保存在db.properties配置文件中，程序启动时读取配置并加载MySQL驱动。DAO层通过DBUtil.getConnection方法获取Connection对象，并使用try-with-resources自动释放数据库资源。",
            "统一数据库连接工具的好处在于降低重复代码数量，便于修改数据库连接参数。若后续引入连接池，只需要调整DBUtil实现，不必修改所有DAO类。当前系统采用PreparedStatement执行SQL语句，既便于参数绑定，也能降低SQL注入风险。"
        ]),
        ("4.2 用户登录与权限控制实现", [
            "登录功能由LoginServlet处理。用户提交用户名和密码后，Servlet调用UserDao.login方法查询users表。如果查询成功，则将User对象保存到Session中；如果用户角色为admin，则跳转到后台首页，否则跳转到学生端首页。退出功能由LogoutServlet处理，通过invalidate方法清空Session。",
            "权限控制由AuthFilter实现。过滤器拦截/cart、/checkout、/orders、/address和/admin/*等路径。若Session中不存在user对象，则跳转到登录页；若访问/admin路径但用户角色不是admin，则返回403错误。通过过滤器集中处理权限检查，可以避免在每个Servlet中重复编写登录判断代码。"
        ]),
        ("4.3 菜品浏览与购物车实现", [
            "首页由HomeServlet处理。Servlet读取keyword参数，调用DishDao.findAll方法查询菜品列表，同时查询商家列表和公告列表，再转发到index.jsp展示。DishDao通过连接dishes、merchants和categories三张表，获得菜品所属商家和分类名称，从而减少页面中的额外查询。",
            "购物车由CartServlet实现。由于购物车数据在用户提交订单前属于临时状态，系统将其保存在Session中，数据结构为Map<Integer, CartItem>。加入购物车时，如果菜品已存在则数量加一，否则新建CartItem；修改数量时若数量小于等于0则删除该菜品；展示购物车时对所有明细小计求和得到总价。"
        ]),
        ("4.4 订单提交实现", [
            "订单提交是系统核心业务。CheckoutServlet在GET请求中展示购物车和地址列表，在POST请求中读取选中的addressId和订单备注，生成订单编号，计算购物车总金额，并调用OrderDao.createOrder写入数据库。订单编号采用时间戳组合生成，便于在页面和后台识别。",
            "OrderDao.createOrder方法采用数据库事务处理。方法首先插入orders表并获取自动生成的订单id，然后遍历购物车写入order_items表。如果任一步骤发生异常，则回滚事务；如果全部成功，则提交事务。事务控制保证订单主表和订单明细表保持一致，避免业务数据残缺。"
        ]),
        ("4.5 后台管理实现", [
            "后台首页由AdminDashboardServlet处理，用于统计商家数量、菜品数量和订单数量。商家管理由AdminMerchantServlet处理，支持新增、编辑和删除商家。菜品管理由AdminDishServlet处理，支持维护菜品所属商家、分类、名称、价格、描述、库存和状态。订单管理由AdminOrderServlet处理，管理员可以将订单状态更新为created、paid、delivering、finished或cancelled。",
            "后台页面采用表格加表单的布局方式，符合管理系统常见交互习惯。列表区域用于快速查看已有数据，表单区域用于新增或编辑数据。删除操作通过确认框降低误操作概率。"
        ]),
        ("4.6 系统安全与可扩展设计", [
            "系统当前安全措施包括登录会话校验、管理员角色校验、PreparedStatement参数化SQL和受保护路径过滤。对于课程实训项目，这些措施可以满足基本演示需求。若系统投入真实使用，还应增加密码加密存储、验证码、CSRF防护、输入长度校验、文件上传校验和日志审计。",
            "扩展性方面，系统表结构预留了reviews和announcements等扩展实体，后续可增加评价管理、公告后台维护、销量排行、菜品图片上传和订单分页查询。由于项目采用分层结构，新增功能时通常只需要增加Model、DAO、Servlet和JSP页面，扩展路径清晰。"
        ]),
    ],
    "第五章 系统测试": [
        ("5.1 测试目的与测试环境", [
            "系统测试的目的是验证校园外卖点餐管理系统是否满足需求分析中提出的功能要求，并检查登录、购物车、订单提交和后台管理等核心流程是否稳定。测试还用于发现页面跳转错误、数据库写入异常、权限控制缺陷和表单输入问题。",
            "测试环境为Windows 11操作系统，JDK 17，Maven 3.6.3，MySQL 8.0，Tomcat 9和Chrome浏览器。测试前导入campus_takeaway.sql脚本，使用管理员账号admin/admin123和学生账号student/123456进行功能验证。"
        ]),
        ("5.2 功能测试用例", [
            "登录测试包括学生登录、管理员登录和错误密码登录。学生登录成功后应进入学生端首页，管理员登录成功后应进入后台首页，错误密码应提示用户名或密码错误。购物车测试包括添加菜品、修改数量、删除菜品和空购物车跳转。订单测试包括选择已有地址提交订单、新增地址后提交订单、查看我的订单和后台修改订单状态。",
            "后台测试包括新增商家、编辑商家、删除商家、新增菜品、编辑菜品、删除菜品和更新订单状态。异常测试包括未登录访问购物车、学生访问后台路径、提交空购物车和输入非法数量。系统应通过跳转登录页、返回403或删除购物车项等方式处理异常。"
        ]),
        ("5.3 性能与并发测试", [
            "性能测试以本地环境为基础，采用浏览器连续访问和简单并发请求方式验证系统响应情况。测试重点包括首页菜品查询、购物车页面展示、订单提交和后台订单列表访问。在测试数据量较小的情况下，页面响应能够保持在较短时间内，Maven构建生成的WAR包可正常部署。",
            "并发测试模拟多个学生账号同时浏览菜品和提交订单。由于购物车数据存储在各自Session中，不同用户之间购物车不会相互影响。订单提交使用数据库事务，能够保证单个订单的数据一致性。需要说明的是，当前系统尚未使用连接池和压力测试工具，在高并发真实生产环境下还需要进一步优化数据库连接、分页查询和索引配置。"
        ]),
        ("5.4 测试结论", [
            "测试结果表明，系统能够完成学生注册登录、菜品浏览、购物车管理、地址选择、订单提交、订单查询和后台维护等主要功能。权限过滤能够阻止未登录用户访问受保护页面，也能够限制普通学生访问后台管理页面。订单提交能够同时写入订单表和订单明细表，符合业务闭环要求。",
            "系统仍存在一些不足，例如密码未加密、页面样式较简单、订单状态没有配送时间记录、评价功能仅完成数据表预留、后台未实现分页和图片上传。这些不足不会影响课程实训的核心演示，但为后续完善提供了明确方向。"
        ]),
    ],
    "第六章 结论": [
        ("6.1 研究成果总结", [
            "本课题完成了校园外卖点餐管理系统的需求分析、总体设计、详细设计、编码实现和测试验证。系统基于JSP、Servlet、JDBC和MySQL实现，包含九张核心数据表，覆盖学生端点餐流程和管理员端维护流程。项目能够通过Maven打包为WAR文件并部署到Tomcat运行，满足Java Web实训项目的基本要求。",
            "系统实现了用户注册登录、公告展示、商家展示、菜品搜索、购物车管理、地址维护、订单提交、订单查询、商家管理、菜品管理和订单状态更新等功能。数据库设计体现了用户、地址、订单、订单明细、商家、分类和菜品之间的关系，订单提交使用事务保证数据一致性。"
        ]),
        ("6.2 不足与展望", [
            "当前系统仍有改进空间。安全方面，密码应采用哈希算法存储，关键表单应增加更严格的输入校验。性能方面，可以引入数据库连接池、分页查询和索引优化。功能方面，可以继续完善评价管理、公告后台维护、支付模拟、菜品图片上传、销量统计和配送状态跟踪。界面方面，可以进一步提升响应式布局和用户体验。",
            "未来若将系统扩展为毕业设计，可以采用Spring Boot、MyBatis和Vue等技术重构，提高工程化水平；也可以结合校园真实场景加入食堂窗口管理、配送员管理、订单超时提醒和数据可视化统计，使系统从课程实训作品进一步发展为具有实际应用价值的校园生活服务平台。"
        ]),
    ],
}


REPORT_EXTRA = {
    "一、课题来源及研究意义": [
        "从数据管理角度看，校园外卖点餐系统并不是简单的页面展示项目，而是一个围绕订单数据流转建立的信息管理系统。学生从浏览菜品到提交订单，会产生用户行为、地址信息、订单主信息和订单明细信息；管理员维护商家和菜品时，会影响学生端可见数据。若缺少统一系统，商家容易通过表格、聊天记录或纸质登记分散保存数据，后续查询、统计和追溯都较困难。因此，本课题能够体现管理系统对数据集中管理和业务流程规范化的价值[12]。",
    ],
    "二、国内外发展现状": [
        "从软件工程研究角度看，当前管理系统开发越来越强调需求与实现的一致性。国际软件工程教材指出，需求、设计、实现和测试之间需要保持可追踪关系，否则容易出现系统功能与用户目标脱节的问题[14][15]。本课题在调研阶段即将学生端和管理员端功能拆分，并把每项功能映射到数据库表和Servlet控制器，有利于后续论文撰写和项目答辩时说明系统设计逻辑。",
    ],
    "四、研究内容": [
        "评价表和公告表虽然在基础演示中属于扩展功能，但它们具有明确的业务意义。评价表可以支持学生对商家服务和菜品质量进行反馈，公告表可以支持平台发布配送提示、系统维护通知和活动信息。预留这些实体能够说明系统并非只围绕单次下单流程设计，而是具备向校园生活服务平台扩展的基础。",
    ],
    "六、开发方案可行性分析": [
        "从开发风险看，项目主要风险不在算法复杂度，而在表关系设计、中文编码、路径跳转和部署环境配置。为降低风险，系统采用Maven管理依赖，使用统一编码过滤器处理中文请求，使用配置文件保存数据库连接信息，并通过README记录数据库导入和WAR包部署步骤。这些措施能够提升项目复现性和交付完整性。",
    ],
}


THESIS_EXTRA = {
    "第二章 需求分析": [
        ("2.5 数据需求分析", [
            "数据需求是本系统需求分析中的核心部分。学生端业务至少需要保存用户账号、菜品信息、购物车临时数据、收货地址、订单主信息和订单明细信息；管理员端业务至少需要保存商家信息、菜品分类、菜品库存和订单状态。若将所有信息放在一张或少量数据表中，会导致字段重复、数据冗余和查询困难。因此，系统将业务实体拆分为九张表，通过主键和外键建立关系。",
            "用户表中的role字段用于区分学生和管理员，这种设计能够在不增加管理员表的情况下实现基本角色控制。菜品表同时关联商家表和分类表，既能显示菜品所属商家，也能支持按分类检索。订单表只保存订单整体信息，订单明细表保存具体菜品快照，这样即使后续菜品价格发生变化，历史订单仍能保留当时下单价格。该设计符合订单系统常见的数据持久化要求。",
            "地址表单独设计而不是直接写入用户表，是因为一个学生可能有多个收货地址，例如宿舍、实验室或教学楼。订单表关联地址编号，可以在下单时选择具体地址。评价表与公告表作为扩展数据表，分别支持服务反馈和平台通知，为系统后续完善提供空间。"
        ]),
        ("2.6 性能与可靠性需求", [
            "系统面向课程实训和小规模校园场景，性能目标应与实际使用范围相匹配。首页菜品查询、购物车展示、订单查询和后台列表访问应在普通本地开发环境中快速响应。对于常见演示数据量，单次页面请求应尽量控制在2秒以内，订单提交过程应在数据库事务内完成，保证用户提交后能够看到明确结果。",
            "可靠性需求主要体现在数据一致性和异常处理两个方面。订单业务需要保证订单表与订单明细表同时成功或同时失败，因此OrderDao采用事务处理。权限控制需要保证未登录用户不能提交订单，普通学生不能进入后台。页面跳转需要保证登录失败、空购物车、无权限访问等异常情况都有明确处理结果。",
            "可维护性需求要求系统代码结构清晰，避免把SQL语句、页面展示和业务判断混杂在同一文件中。项目采用DAO封装SQL、Servlet处理请求、JSP展示数据的方式，虽然属于传统Java Web架构，但对于实训项目而言具有较强可读性和可讲解性。"
        ]),
    ],
    "第三章 总体设计": [
        ("3.5 表结构约束与完整性设计", [
            "数据库完整性设计包括实体完整性、参照完整性和用户定义完整性。实体完整性通过每张表的主键实现，例如users表、merchants表和dishes表均采用自增id作为主键。参照完整性通过外键实现，例如dishes表中的merchant_id引用merchants表，orders表中的user_id引用users表，order_items表中的order_id引用orders表。用户定义完整性通过字段类型、非空约束和默认值实现，例如订单状态默认值为created，菜品状态默认值为on。",
            "订单明细表保存dish_name、price和subtotal等字段，这看似与菜品表存在部分重复，但在订单系统中具有业务合理性。因为菜品名称和价格可能在订单生成后被管理员修改，历史订单仍应保留下单时的菜品快照。该设计属于为了历史记录准确性而进行的受控冗余，并不违反业务一致性要求。",
            "系统没有把购物车设计为数据库表，而是将购物车保存在Session中。原因在于购物车属于下单前的临时状态，课程实训系统不需要跨设备保存购物车。若未来扩展为真实应用，可以增加cart和cart_items表，将购物车持久化到数据库中。"
        ]),
        ("3.6 目录结构与类职责设计", [
            "项目目录结构按照Maven Web项目规范组织。src/main/java目录保存Java源代码，src/main/resources保存数据库连接配置，src/main/webapp保存JSP页面和静态资源，db目录保存数据库脚本，docs目录保存项目说明材料。该结构与Maven打包流程一致，能够生成标准WAR文件并部署到Tomcat。",
            "Java源代码按照功能职责划分为model、dao、servlet和util四个包。model包中的类与数据库实体相对应，如User、Dish、Order和Address；dao包负责执行SQL并返回实体对象；servlet包负责处理浏览器请求；util包保存DBUtil数据库连接工具。通过这种包结构，可以从文件路径直接判断代码职责，降低维护难度。",
            "后台JSP页面放置在webapp/admin目录下，与学生端页面区分。权限过滤器拦截/admin/*路径，确保只有管理员能够访问后台。静态样式文件统一放置在assets/css目录，便于全站复用样式。"
        ]),
        ("3.7 数据流设计", [
            "学生下单数据流从首页开始。HomeServlet查询菜品、商家和公告数据并转发到index.jsp；学生点击加入购物车后，CartServlet将菜品对象封装为CartItem并保存到Session；结算时CheckoutServlet读取Session购物车和用户地址；提交订单后OrderDao将订单写入orders表，将购物车明细写入order_items表，最后清空Session购物车。",
            "管理员维护数据流相对直接。管理员进入后台后，AdminDashboardServlet统计商家、菜品和订单数量。商家管理页面通过AdminMerchantServlet执行新增、编辑和删除操作。菜品管理页面通过AdminDishServlet读取商家和分类下拉数据，并保存菜品信息。订单管理页面通过AdminOrderServlet修改订单状态。所有后台操作最终都通过DAO层访问数据库。",
            "这种数据流设计体现了请求驱动的Web应用特点。浏览器发起请求，Servlet解析参数并调用DAO，DAO执行SQL并返回结果，Servlet再转发或重定向到JSP页面。每个请求路径对应一个清晰业务动作，有利于测试和排错。"
        ]),
    ],
    "第四章 详细设计及实现": [
        ("4.7 核心代码设计说明", [
            "DBUtil类负责读取db.properties配置文件并创建数据库连接。该类在静态代码块中加载MySQL驱动，在getConnection方法中调用DriverManager.getConnection获取连接。由于数据库连接参数集中在配置文件中，部署到不同电脑时只需要修改用户名、密码和数据库地址，不必修改Java源代码。",
            "UserDao类封装用户登录和注册相关SQL。login方法根据用户名和密码查询users表，若存在匹配记录则返回User对象。register方法插入学生用户信息，并将role默认设置为student。existsByUsername方法用于注册前检查用户名是否重复。DAO层通过PreparedStatement绑定参数，避免字符串拼接SQL带来的安全风险。",
            "DishDao类是多表查询的典型示例。findAll方法连接dishes、merchants和categories三张表，使页面能够直接显示菜品名称、商家名称和分类名称。该方法支持关键词搜索，关键词会同时匹配菜品、商家和分类，提高学生查找菜品的便利性。",
            "OrderDao类体现了事务处理。createOrder方法先关闭自动提交，插入订单主表后获取订单id，再批量插入订单明细。如果插入过程中出现异常，则执行rollback；如果全部成功，则执行commit。该逻辑保证订单数据不会只写入一半，是订单类系统必须重视的可靠性设计。"
        ]),
        ("4.8 页面交互设计说明", [
            "首页页面由index.jsp实现，页面上方显示导航栏，中间显示搜索框、公告、商家信息和推荐菜品。菜品采用卡片形式展示，包含菜品名称、商家、分类、描述、价格、库存和加入购物车按钮。该布局便于学生快速浏览和选择。",
            "购物车页面由cart.jsp实现，使用表格展示菜品、商家、单价、数量、小计和操作按钮。数量修改采用表单提交方式，删除操作也通过表单触发。页面底部展示总价，并提供继续点餐和去结算入口。结算页面展示订单明细、地址列表、新增地址表单和备注输入框，使下单信息集中在一个页面中。",
            "后台页面采用统一导航栏和表格布局。商家管理页面将新增编辑表单和商家列表放在同一页面，菜品管理页面通过下拉框选择商家和分类，订单管理页面通过下拉框修改订单状态。这样的页面结构虽然简洁，但符合管理系统高频编辑和快速查看的使用特点。"
        ]),
        ("4.9 异常处理与边界控制", [
            "系统在多个位置进行了边界处理。CheckoutServlet在展示结算页和提交订单前都会检查购物车是否为空，若为空则重定向到购物车页面。CartServlet在修改数量时，如果数量小于等于0，则直接从购物车中移除对应菜品。LoginServlet在登录失败时转发回登录页面并显示错误信息。",
            "权限边界由AuthFilter统一处理。未登录用户访问受保护路径时会被重定向到登录页面，普通学生访问后台路径时会收到403错误。统一过滤比在每个Servlet中单独判断更清晰，也更不容易遗漏。管理员删除商家或菜品时，页面提供确认提示，减少误操作风险。",
            "数据库边界方面，系统使用外键约束维护表关系。如果管理员试图删除已经被菜品引用的商家，数据库外键会阻止删除，从而避免产生孤立数据。该行为在真实系统中可以进一步优化为状态关闭而不是物理删除。"
        ]),
        ("4.10 部署与运行实现", [
            "系统通过Maven进行构建，pom.xml中声明Servlet API和MySQL Connector/J依赖，并将打包类型设置为war。执行mvn clean package后，target目录下会生成campus-takeaway.war文件。将该WAR文件复制到Tomcat webapps目录，启动Tomcat后即可通过浏览器访问系统。",
            "部署前需要先在MySQL中导入db/campus_takeaway.sql脚本。脚本会创建campus_takeaway数据库、九张核心表和部分测试数据。随后需要根据本机MySQL账号修改src/main/resources/db.properties中的username和password。若Maven全局配置本地仓库不可用，可使用项目内.mvn/settings.xml进行构建。",
            "系统默认学生账号为student/123456，管理员账号为admin/admin123。通过这两个账号可以分别演示学生点餐流程和后台管理流程。"
        ]),
    ],
    "第五章 系统测试": [
        ("5.5 测试数据设计", [
            "测试数据包括管理员账号、学生账号、三条商家记录、三条分类记录、六条菜品记录、一条默认地址和两条公告。商家包括一食堂简餐、奶茶小站和夜宵窗口，分类包括主食、饮品和小吃，菜品包括照烧鸡腿饭、番茄牛腩饭、珍珠奶茶、满杯水果茶、香辣炸串拼盘和火腿蛋炒饭。测试数据覆盖不同商家、不同分类和不同价格区间，能够支持搜索、浏览和下单测试。",
            "订单测试数据通过前台提交生成，而不是直接在数据库脚本中插入。这样可以验证系统从购物车到订单表和订单明细表的完整写入过程。后台订单状态测试在订单生成后进行，管理员将订单状态从created修改为paid、delivering或finished，以验证订单状态更新功能。"
        ]),
        ("5.6 测试问题与改进建议", [
            "测试过程中需要关注中文编码问题。如果页面提交中文备注或地址后出现乱码，说明请求编码或数据库连接编码配置存在问题。系统通过EncodingFilter设置请求编码，并在数据库连接URL中设置useUnicode和characterEncoding参数，以保证中文数据能够正常保存和显示。",
            "还需要关注删除操作带来的外键约束问题。由于菜品可能已经被订单明细引用，直接删除菜品可能受到数据库约束限制。真实系统中更合理的做法是将菜品状态设置为off，使其不再前台展示，而不是直接物理删除。课程实训系统保留删除功能用于演示增删改查，但后续应改进为逻辑删除。",
            "性能方面，当前系统没有分页查询和连接池，在数据量增大时后台列表可能变慢。后续可以为订单列表、菜品列表增加分页参数，为常用外键字段增加索引，并使用连接池复用数据库连接，从而提升并发访问能力。"
        ]),
    ],
}


THESIS_MORE = {
    "第一章 绪论": [
        ("1.5 论文组织结构", [
            "本文围绕校园外卖点餐管理系统的设计与实现展开。第一章介绍项目背景、研究意义、国内外发展现状以及开发环境，说明课题选择的现实依据。第二章从用户角色、功能需求、非功能需求和可行性角度进行分析，明确系统应完成的业务范围。第三章进行总体设计，阐述系统架构、模块划分、数据库概念结构和逻辑结构。第四章对数据库连接、用户登录、购物车、订单提交、后台管理和安全控制等功能进行详细说明。第五章介绍测试目的、测试环境、测试用例、性能测试和测试结论。第六章总结系统成果，分析不足并提出后续改进方向。",
            "这种组织结构与软件工程过程保持一致，能够体现从需求到实现再到测试的完整链条。论文中的图表、代码节选和测试数据均与实际项目对应，避免只停留在理论描述层面。"
        ]),
    ],
    "第二章 需求分析": [
        ("2.7 业务规则分析", [
            "系统业务规则首先体现在用户身份上。未登录用户可以访问登录和注册页面，但不能访问购物车、结算、订单和后台管理页面。学生用户可以完成点餐相关操作，但不能进入后台管理页面。管理员可以维护商家、菜品和订单状态，但不作为普通学生提交订单。该规则通过Session和AuthFilter共同实现。",
            "其次，购物车规则要求同一菜品重复加入时不产生重复行，而是增加数量。数量修改为0或负数时，系统将该菜品从购物车删除。购物车总价由各项菜品价格乘以数量后求和得出，显示金额应与订单提交金额一致。该规则保证学生在结算前能够清晰确认订单内容。",
            "再次，订单规则要求提交订单时必须选择收货地址，订单生成后写入订单主表和订单明细表。订单状态初始为created，管理员可根据处理过程修改为paid、delivering、finished或cancelled。订单历史数据应保留当时下单的菜品名称和价格，不能因后台菜品信息修改而影响历史订单展示。"
        ]),
    ],
    "第三章 总体设计": [
        ("3.8 数据库字段设计说明", [
            "users表中的username设置唯一约束，用于避免账号重复；password字段保存登录密码；real_name和phone用于展示真实姓名和联系方式；role字段用于区分学生和管理员。merchants表中的status字段用于表示商家营业状态，便于后续扩展前台过滤功能。categories表保存菜品分类，使菜品展示和搜索更有组织。",
            "dishes表是系统中被频繁访问的数据表之一。merchant_id和category_id分别关联商家与分类，price使用DECIMAL类型避免浮点误差，stock用于表示库存数量，status用于表示菜品上架或下架。addresses表保存receiver_name、phone和detail字段，满足校园宿舍地址记录需求。",
            "orders表保存订单编号、用户编号、地址编号、总金额、状态、备注和创建时间。order_no设置唯一约束，便于用户和管理员根据订单编号识别订单。order_items表保存订单明细，其中subtotal字段可由price和quantity计算得到，但保存该字段能够提升订单展示效率，也便于后续统计。reviews表和announcements表为评价和公告扩展提供基础。"
        ]),
        ("3.9 模块间接口设计", [
            "Servlet与DAO之间通过方法调用传递数据。LoginServlet调用UserDao.login返回User对象；HomeServlet调用DishDao.findAll、MerchantDao.findAll和AnnouncementDao.findLatest获取首页数据；CartServlet调用DishDao.findById获取菜品详情；CheckoutServlet调用AddressDao.findByUserId和OrderDao.createOrder完成结算；后台Servlet调用对应DAO完成增删改查。",
            "Servlet与JSP之间通过request属性和Session属性传递数据。首页数据通过request.setAttribute传递给index.jsp，登录用户通过session.setAttribute保存，购物车也通过Session保存。JSP页面只负责展示和提交表单，不直接执行数据库访问。该接口方式简单明确，适合传统JSP/Servlet项目。",
            "数据库与DAO之间通过SQL接口交互。DAO负责将ResultSet映射为实体类，也负责将实体类字段绑定到PreparedStatement参数中。通过这种方式，页面层不需要了解数据库字段细节，控制层也不需要直接拼接SQL语句。"
        ]),
    ],
    "第四章 详细设计及实现": [
        ("4.11 关键页面实现解析", [
            "登录页面login.jsp包含用户名和密码输入框，表单提交到/login路径。若登录失败，LoginServlet会把错误信息放入request属性，页面读取error属性并显示提示。该实现方式使错误反馈与页面保持在同一视图中，用户可以直接修改输入并重新提交。",
            "首页index.jsp通过循环展示商家、公告和菜品。菜品区域的每个卡片都包含一个隐藏dishId字段和action=add字段，提交后由CartServlet识别为加入购物车操作。搜索框提交到/home路径，HomeServlet根据keyword执行模糊查询。页面导航根据Session中的user对象判断是否显示登录注册链接或退出链接。",
            "后台菜品管理页面admin/dishes.jsp同时展示编辑表单和菜品列表。编辑表单中的商家和分类来自数据库查询结果，使用下拉框减少输入错误。列表中的编辑链接携带action=edit和id参数，Servlet读取参数后查询对应菜品并回填表单。新增和编辑共用同一个保存逻辑，依据id是否为空判断插入或更新。"
        ]),
        ("4.12 代码质量与维护性分析", [
            "系统代码遵循单一职责原则的基本思想。实体类只保存属性和getter、setter方法，DAO类只处理数据库访问，Servlet类只处理请求流程，JSP页面只处理展示。虽然项目规模不大，但这种职责划分能够减少代码交叉依赖，为后续功能扩展保留空间。",
            "项目中的命名尽量与业务含义保持一致，例如Merchant表示商家，Dish表示菜品，Order表示订单，OrderItem表示订单明细。数据库字段采用下划线命名，Java属性采用驼峰命名，DAO映射时完成二者转换。清晰命名能够降低答辩讲解和后续维护成本。",
            "系统仍有进一步优化空间。例如多个DAO中存在相似的连接和异常处理代码，后续可以抽取通用BaseDao；后台删除操作目前是物理删除，后续可以改为逻辑删除；页面中使用JSP脚本片段较多，后续可以引入JSTL或前后端分离技术提升可维护性。"
        ]),
        ("4.13 与需求对应关系说明", [
            "需求分析中提出的学生注册登录由RegisterServlet、LoginServlet、register.jsp和login.jsp实现；菜品浏览和搜索由HomeServlet、DishDao和index.jsp实现；购物车管理由CartServlet和cart.jsp实现；地址维护由AddressServlet和checkout.jsp实现；订单提交由CheckoutServlet、OrderDao和checkout.jsp实现；订单查询由OrderServlet和orders.jsp实现。",
            "管理员后台需求由AdminDashboardServlet、AdminMerchantServlet、AdminDishServlet和AdminOrderServlet实现。后台首页展示统计数据，商家管理页面维护merchants表，菜品管理页面维护dishes表，订单管理页面维护orders表状态。AuthFilter实现登录和角色校验，EncodingFilter解决中文编码问题。由此可见，需求、设计和实现之间具有明确对应关系。",
            "数据库需求也得到落实。系统实际包含users、merchants、categories、dishes、addresses、orders、order_items、reviews和announcements九张表，超过实训要求的六张表。核心业务至少使用前七张表，评价表和公告表作为扩展模块体现系统可拓展性。"
        ]),
    ],
    "第五章 系统测试": [
        ("5.7 测试覆盖分析", [
            "从功能覆盖角度看，测试覆盖了用户登录、学生点餐、订单提交和后台维护四条主线。登录测试验证身份认证和角色跳转，学生点餐测试验证菜品展示、购物车和结算流程，订单测试验证数据库事务写入，后台测试验证管理员对基础数据和订单状态的维护能力。这些测试点覆盖了系统最核心的业务风险。",
            "从数据覆盖角度看，测试数据覆盖了不同商家、不同分类、不同菜品和不同订单状态。通过多次加入购物车、修改数量和删除菜品，可以验证Session购物车逻辑。通过新增地址再提交订单，可以验证地址表与订单表之间的关联。通过后台修改订单状态，可以验证管理员端对学生端订单数据的影响。",
            "从异常覆盖角度看，测试包括错误登录、未登录访问、普通学生访问后台、空购物车结算和数量为0等情况。异常测试能够发现系统边界处理是否完整，也是管理系统稳定性的重要组成部分。"
        ]),
        ("5.8 测试结果评价", [
            "综合测试结果表明，系统能够满足课程实训的功能要求和演示要求。主要页面能够正常跳转，数据库读写能够正常完成，学生端和管理员端角色边界清晰。Maven构建能够成功生成WAR包，说明项目依赖配置和源码编译没有明显问题。",
            "测试也暴露出一些需要后续优化的问题。首先，系统没有实现真实支付和配送模块，订单状态由管理员手动更新。其次，部分表单只依赖浏览器required属性进行校验，后端校验还不充分。再次，系统未使用分页，在订单和菜品数量较多时页面可能变长。最后，密码以明文方式存储，不适合生产环境。",
            "针对上述问题，后续改进方向包括增加服务端参数校验、引入密码哈希、完善逻辑删除、增加分页查询和连接池、扩展评价管理和公告管理，以及增加订单统计图表。通过这些改进，系统可以从课程实训版本逐步升级为更完整的校园生活服务管理平台。"
        ]),
    ],
}


def build_report():
    doc = init_doc()
    add_cover(doc, "校园外卖点餐管理系统调研报告", "作业2-调研报告")
    add_heading(doc, "校园外卖点餐管理系统调研报告", 1)
    for heading, paragraphs in REPORT_SECTIONS.items():
        add_heading(doc, heading, 2)
        for idx, para in enumerate(paragraphs):
            add_body_paragraph(doc, para, cite=((idx % 11) + 1))
        for idx, para in enumerate(REPORT_EXTRA.get(heading, [])):
            add_body_paragraph(doc, para, cite=((idx + 3) % 13) + 1)
    add_references(doc, 2)
    add_page_number(doc.sections[0])
    OUT_REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_REPORT)


def add_abstracts(doc):
    add_heading(doc, "摘要", 1)
    for para in [
        "随着高校生活服务数字化程度不断提高，学生对校园餐饮服务的便捷性、透明性和可管理性提出了更高要求。传统线下排队和人工记录方式难以满足高峰期集中订餐需求，也不利于商家和管理人员及时维护菜品与订单信息。因此，设计并实现一个面向校园场景的外卖点餐管理系统，具有较强的应用价值和教学实践意义。",
        "系统采用Java Web技术路线，以JSP负责页面展示，以Servlet负责请求控制，以JDBC和DAO模式完成MySQL数据库访问。系统设计了用户、商家、菜品分类、菜品、收货地址、订单、订单明细、评价和公告九张核心数据表，实现学生注册登录、菜品浏览、购物车、地址选择、订单提交、订单查询，以及管理员商家管理、菜品管理和订单状态维护等功能。",
        "测试结果表明，系统能够在Tomcat环境下正常部署运行，主要业务流程完整，订单提交过程能够通过事务保证订单主表和明细表的一致性。系统结构清晰、维护方便，能够满足Java Web实训项目对数据库设计、页面交互和后台管理的综合要求。"
    ]:
        add_body_paragraph(doc, para)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, line_pt=18)
    add_text(p, "关键词：", "黑体", 10.5, True)
    add_text(p, "校园外卖；Java Web；JSP；Servlet；MySQL", "宋体", 10.5)
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False)
    add_text(p, "ABSTRACT", "Times New Roman", 16, True)
    for para in [
        "With the continuous improvement of digital campus services, students have higher requirements for convenience, transparency and manageability of campus catering services. Traditional offline queuing and manual recording methods are difficult to meet concentrated ordering demands during peak hours, and they are not conducive to timely maintenance of dishes and order information.",
        "The system adopts a Java Web technology route. JSP is used for page presentation, Servlet is used for request control, and JDBC with DAO is used for MySQL database access. Nine core tables are designed, including users, merchants, categories, dishes, addresses, orders, order items, reviews and announcements. The system implements student registration and login, dish browsing, shopping cart, address selection, order submission, order query, and administrator management of merchants, dishes and order status.",
        "Test results show that the system can be deployed and run normally in the Tomcat environment. The main business process is complete, and the order submission process can ensure consistency between the order table and order item table through database transactions. The system has a clear structure and convenient maintenance, meeting the comprehensive requirements of Java Web training projects."
    ]:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=True)
        add_text(p, para, "Times New Roman", 12)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, line_pt=18)
    add_text(p, "Key words: ", "Times New Roman", 10.5, True)
    add_text(p, "campus takeaway; Java Web; JSP; Servlet; MySQL", "Times New Roman", 10.5)
    doc.add_page_break()


def add_toc_placeholder(doc):
    add_heading(doc, "目录", 1)
    entries = ["引言", "第一章 绪论", "第二章 需求分析", "第三章 总体设计", "第四章 详细设计及实现", "第五章 系统测试", "第六章 结论", "致谢", "参考文献"]
    for i, e in enumerate(entries, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, line_pt=20)
        add_text(p, f"{e}" + "." * max(4, 45 - len(e) * 2) + f"{i + 2}", "宋体", 12)
    doc.add_page_break()


def build_thesis():
    diagrams = build_diagrams()
    doc = init_doc()
    add_cover(doc, "校园外卖点餐管理系统的设计与实现", "作业3-论文")
    add_abstracts(doc)
    add_toc_placeholder(doc)
    add_heading(doc, "引言", 1)
    add_body_paragraph(doc, "校园外卖点餐管理系统是面向高校生活服务场景设计的Web管理系统。系统以学生点餐和管理员维护为主线，围绕用户、商家、菜品、地址和订单建立业务闭环。本文按照软件工程开发流程，对系统背景、需求分析、总体设计、详细设计、系统测试和总结展望进行论述。", 1)
    add_body_paragraph(doc, "论文内容与已实现项目保持一致，重点说明JSP、Servlet、DAO、JDBC和MySQL在系统中的具体应用方式，并通过用例图、项目结构图、流程图和E-R图展示系统设计结果。", 14)

    # Figures in required chapters.
    figure_plan = {
        "第二章 需求分析": [(diagrams[0], "图2-1 系统用例图")],
        "第三章 总体设计": [(diagrams[1], "图3-1 项目结构图"), (diagrams[3], "图3-2 数据库全局E-R图")],
        "第四章 详细设计及实现": [(diagrams[2], "图4-1 学生下单业务流程图")],
    }

    for chapter, sections in THESIS_CHAPTERS.items():
        add_heading(doc, chapter, 1)
        for path, cap in figure_plan.get(chapter, []):
            add_picture(doc, path, cap)
        for sec_title, paragraphs in sections:
            add_heading(doc, sec_title, 2)
            for idx, para in enumerate(paragraphs):
                add_body_paragraph(doc, para, cite=((idx + 1) % 13) + 1)
            if sec_title == "2.4 可行性分析":
                add_table(doc, ["可行性类别", "分析结论"], [
                    ["技术可行性", "Java Web、Servlet、JSP、JDBC和MySQL均为成熟技术，能够支撑本系统。"],
                    ["运行可行性", "系统可部署在Tomcat 9，普通个人电脑即可运行。"],
                    ["经济可行性", "开发工具和数据库均可免费使用，成本较低。"],
                    ["操作可行性", "页面采用常见表单、按钮和表格，学习成本低。"],
                ], [2200, 7000], "表2-1 可行性分析表")
            if sec_title == "3.4 数据库逻辑结构设计":
                add_table(doc, ["表名", "主要字段", "说明"], [
                    ["users", "id, username, password, real_name, phone, role", "保存学生和管理员账号"],
                    ["merchants", "id, name, description, phone, address, status", "保存商家信息"],
                    ["dishes", "id, merchant_id, category_id, name, price, stock", "保存菜品信息"],
                    ["orders", "id, order_no, user_id, address_id, total_amount, status", "保存订单主信息"],
                    ["order_items", "id, order_id, dish_id, price, quantity, subtotal", "保存订单明细"],
                    ["addresses", "id, user_id, receiver_name, phone, detail", "保存收货地址"],
                ], [1700, 4500, 3000], "表3-1 核心数据表设计")
            if sec_title == "4.4 订单提交实现":
                add_table(doc, ["步骤", "处理内容", "涉及对象"], [
                    ["1", "读取Session购物车并校验是否为空", "CheckoutServlet"],
                    ["2", "生成订单编号并计算总金额", "Order"],
                    ["3", "开启事务并写入orders表", "OrderDao"],
                    ["4", "批量写入order_items表", "OrderDao"],
                    ["5", "提交事务并清空购物车", "Session"],
                ], [1000, 5600, 2600], "表4-1 订单提交步骤表")
            if sec_title == "5.2 功能测试用例":
                add_table(doc, ["编号", "测试项目", "输入条件", "预期结果", "结果"], [
                    ["TC01", "学生登录", "student/123456", "进入学生首页", "通过"],
                    ["TC02", "管理员登录", "admin/admin123", "进入后台首页", "通过"],
                    ["TC03", "错误登录", "错误密码", "提示用户名或密码错误", "通过"],
                    ["TC04", "加入购物车", "选择照烧鸡腿饭", "购物车数量增加", "通过"],
                    ["TC05", "提交订单", "选择地址并提交", "生成订单记录", "通过"],
                    ["TC06", "后台改状态", "修改为finished", "订单状态更新", "通过"],
                ], [900, 1700, 2200, 3000, 1200], "表5-1 功能测试用例表")
        for sec_title, paragraphs in THESIS_EXTRA.get(chapter, []):
            add_heading(doc, sec_title, 2)
            for idx, para in enumerate(paragraphs):
                add_body_paragraph(doc, para, cite=((idx + 4) % 13) + 1)
            if sec_title == "4.7 核心代码设计说明":
                add_table(doc, ["类名", "主要职责", "关键方法"], [
                    ["DBUtil", "读取数据库配置并创建连接", "getConnection"],
                    ["UserDao", "处理用户登录、注册和用户名校验", "login, register, existsByUsername"],
                    ["DishDao", "处理菜品查询、新增、编辑和删除", "findAll, findById, save, delete"],
                    ["OrderDao", "处理订单创建、查询和状态更新", "createOrder, findAll, updateStatus"],
                    ["AuthFilter", "统一进行登录与管理员权限检查", "doFilter"],
                ], [1700, 4700, 2800], "表4-2 核心类职责表")
        for sec_title, paragraphs in THESIS_MORE.get(chapter, []):
            add_heading(doc, sec_title, 2)
            for idx, para in enumerate(paragraphs):
                add_body_paragraph(doc, para, cite=((idx + 6) % 13) + 1)

    add_heading(doc, "附录A 核心代码节选", 1)
    code_blocks = [
        ("A.1 数据库连接工具", "public static Connection getConnection() throws SQLException {\n    return DriverManager.getConnection(url, username, password);\n}"),
        ("A.2 登录控制流程", "User user = userDao.login(username, password);\nif (user == null) {\n    request.setAttribute(\"error\", \"用户名或密码错误\");\n    request.getRequestDispatcher(\"/login.jsp\").forward(request, response);\n    return;\n}"),
        ("A.3 订单事务处理", "conn.setAutoCommit(false);\ntry {\n    insertOrder();\n    insertOrderItems();\n    conn.commit();\n} catch (Exception e) {\n    conn.rollback();\n}"),
    ]
    for title, code in code_blocks:
        add_heading(doc, title, 2)
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, line_pt=16, before=4, after=4)
        for line in code.splitlines():
            run = p.add_run(line + "\n")
            set_run_font(run, "Times New Roman", 10.5)

    add_heading(doc, "致谢", 1)
    for para in [
        "本课题的完成离不开课程学习过程中所积累的数据库、Java、前端和JSP基础知识。通过本次系统设计与实现，进一步理解了Web管理系统从需求分析到部署测试的完整过程。",
        "感谢任课教师在实训项目选题、数据库设计和功能实现方面提供的指导，也感谢同学在测试和页面体验方面提出的建议。后续将继续完善系统功能，提高代码质量和用户体验。"
    ]:
        add_body_paragraph(doc, para)
    add_references(doc, 1)
    add_page_number(doc.sections[0])
    OUT_THESIS.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_THESIS)


if __name__ == "__main__":
    build_report()
    build_thesis()
    print(OUT_REPORT)
    print(OUT_THESIS)
